"""DOCX read/write, PDF ingestion, and PDF export."""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn as _qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from rich.console import Console

from .models import (
    Education,
    ResumeDocument,
    ResumeExperience,
    ResumeProject,
    ResumeSection,
    TailoredResume,
)

if TYPE_CHECKING:
    from anthropic import Anthropic

console = Console()

SECTION_KEYWORDS = {
    "summary": ["summary", "objective", "profile", "about"],
    "experience": ["experience", "employment", "work history", "professional experience"],
    "skills": ["skills", "technical skills", "core competencies", "technologies"],
    "education": ["education", "academic", "qualifications"],
    "projects": ["projects", "personal projects", "side projects", "portfolio"],
}

# Resolve the LibreOffice binary once at import time
def _find_soffice() -> Optional[str]:
    for candidate in [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "soffice",
        "libreoffice",
    ]:
        if shutil.which(candidate) or Path(candidate).exists():
            return candidate
    return None

_SOFFICE_BIN: Optional[str] = _find_soffice()


def _detect_section(heading: str) -> Optional[str]:
    h = heading.lower().strip()
    for section, keywords in SECTION_KEYWORDS.items():
        if any(kw in h for kw in keywords):
            return section
    return None


# ── PDF / DOCX reading ────────────────────────────────────────────────────────

def read_resume(path: str | Path, client: Optional["Anthropic"] = None) -> ResumeDocument:
    """Read a PDF or DOCX resume, detecting type by magic bytes."""
    p = Path(path)
    with open(p, "rb") as f:
        if f.read(4) == b"%PDF":
            return read_pdf(p, client)
    return read_docx(p)


def read_pdf(path: str | Path, client: Optional["Anthropic"] = None) -> ResumeDocument:
    """Extract text from PDF and parse into ResumeDocument via Claude."""
    import pdfplumber
    from .claude_engine import _strip_fence

    text_parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=True)
            if page_text:
                text_parts.append(page_text)

    raw_text = "\n".join(text_parts).strip()
    if not raw_text:
        raise ValueError(f"Could not extract text from PDF: {path}")

    if client is None:
        from anthropic import Anthropic as _Anthropic
        from .config import get_anthropic_api_key
        client = _Anthropic(api_key=get_anthropic_api_key())

    console.print("[blue]Parsing PDF resume with Claude...[/blue]")
    from .claude_engine import _claude_create
    response = _claude_create(client,
        model="claude-sonnet-4-6",
        max_tokens=3000,
        messages=[{
            "role": "user",
            "content": f"""Parse this resume text into structured JSON. Return ONLY valid JSON.

Schema:
{{
  "name": "Full Name",
  "contact": "email | phone | location | linkedin",
  "summary": "summary/objective text or null",
  "experience": [
    {{"company": "...", "title": "...", "dates": "...", "bullets": ["..."]}}
  ],
  "skills": ["skill1", "skill2"],
  "education": [
    {{"institution": "...", "degree": "...", "field": "...", "year": "..."}}
  ],
  "projects": [
    {{"name": "...", "description": "...", "tech_stack": ["..."], "bullets": ["..."], "url": null}}
  ]
}}

Resume text:
{raw_text}""",
        }],
    )
    return ResumeDocument.model_validate(json.loads(_strip_fence(response.content[0].text)))


def read_docx(path: str | Path) -> ResumeDocument:
    """Parse a DOCX into a ResumeDocument."""
    doc = Document(str(path))
    raw_sections: list[ResumeSection] = []
    current_heading: Optional[str] = None
    current_lines: list[str] = []

    def flush():
        if current_heading is not None:
            raw_sections.append(
                ResumeSection(heading=current_heading, content="\n".join(current_lines).strip())
            )

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            flush()
            current_heading = text
            current_lines = []
        else:
            if current_heading is None:
                current_heading = "__header__"
            current_lines.append(text)

    flush()

    resume = ResumeDocument(raw_sections=raw_sections)
    for section in raw_sections:
        kind = _detect_section(section.heading)
        if section.heading == "__header__":
            lines = section.content.splitlines()
            if lines:
                resume.name = lines[0]
                resume.contact = " | ".join(lines[1:])
        elif kind == "summary":
            resume.summary = section.content
        elif kind == "skills":
            skills_text = section.content.replace("\n", ", ")
            resume.skills = [s.strip() for s in re.split(r"[,|•·]", skills_text) if s.strip()]
        elif kind == "education":
            resume.education = _parse_education(section.content)
        elif kind == "experience":
            resume.experience = _parse_experience(section.content)
        elif kind == "projects":
            resume.projects = _parse_projects(section.content)

    return resume


def _parse_experience(text: str) -> list[ResumeExperience]:
    experiences: list[ResumeExperience] = []
    current: Optional[dict] = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("-", "•", "·", "*", "–", "—")):
            if current:
                current["bullets"].append(line.lstrip("-•·*–— ").strip())
        else:
            if current:
                experiences.append(ResumeExperience(**current))
            parts = re.split(r"\s*[|–—·]\s*", line)
            if len(parts) >= 2:
                current = {
                    "company": parts[1],
                    "title": parts[0],
                    "dates": parts[2] if len(parts) > 2 else None,
                    "bullets": [],
                }
            else:
                current = {"company": line, "title": "", "dates": None, "bullets": []}

    if current:
        experiences.append(ResumeExperience(**current))
    return experiences


def _parse_education(text: str) -> list[Education]:
    edus: list[Education] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        parts = re.split(r"\s*[|,–—]\s*", line)
        edus.append(Education(
            institution=parts[0] if parts else line,
            degree=parts[1] if len(parts) > 1 else None,
            field=parts[2] if len(parts) > 2 else None,
            year=parts[3] if len(parts) > 3 else None,
        ))
    return edus


def _parse_projects(text: str) -> list[ResumeProject]:
    projects: list[ResumeProject] = []
    current: Optional[dict] = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("-", "•", "·", "*", "–", "—")):
            if current:
                current["bullets"].append(line.lstrip("-•·*–— ").strip())
        else:
            if current:
                projects.append(ResumeProject(**current))
            current = {"name": line, "description": None, "tech_stack": [], "bullets": [], "url": None}

    if current:
        projects.append(ResumeProject(**current))
    return projects


# ── Page count ────────────────────────────────────────────────────────────────

def _get_page_count(docx_path: Path) -> Optional[int]:
    """Convert DOCX to a temp PDF and return page count. Returns None if LibreOffice unavailable."""
    if not _SOFFICE_BIN:
        return None
    tmp_dir = Path(tempfile.mkdtemp())
    try:
        result = subprocess.run(
            [_SOFFICE_BIN, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_dir), str(docx_path)],
            capture_output=True, timeout=30,
        )
        if result.returncode == 0:
            pdf_path = tmp_dir / (docx_path.stem + ".pdf")
            if pdf_path.exists():
                import pdfplumber
                with pdfplumber.open(str(pdf_path)) as pdf:
                    return len(pdf.pages)
    except Exception:
        pass
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    return None


# ── DOCX building ─────────────────────────────────────────────────────────────

def build_docx_from_resume(resume: ResumeDocument, output_path: str | Path) -> None:
    """Build a clean ATS-friendly DOCX, compacting to fit exactly 1 page."""
    output_path = Path(output_path)

    # Fitting schedule: (margin_pt, body_pt, name_pt, space_pt)
    fit_schedule = [
        (36, 10.5, 14, 3),
        (32, 10.0, 13, 2),
        (28,  9.5, 12, 1),
        (24,  9.0, 11, 0),
        (20,  8.5, 10, 0),
    ]

    for margin_pt, body_pt, name_pt, space_pt in fit_schedule:
        _build_docx_raw(resume, output_path, margin_pt, body_pt, name_pt, space_pt)
        pages = _get_page_count(output_path)
        if pages is None:
            console.print("[dim]LibreOffice not found — skipping page-count check.[/dim]")
            break
        if pages <= 1:
            console.print(f"[green]Resume fits in 1 page (font {body_pt}pt, margin {margin_pt}pt).[/green]")
            break
        console.print(f"[yellow]{pages} pages — compacting (→ {body_pt - 0.5}pt)...[/yellow]")
    else:
        console.print("[yellow]Warning: Could not fit resume to 1 page even at minimum size.[/yellow]")


def _build_docx_raw(
    resume: ResumeDocument,
    output_path: Path,
    margin_pt: float,
    body_pt: float,
    name_pt: float,
    space_pt: float,
) -> None:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Pt(margin_pt)
        sec.bottom_margin = Pt(margin_pt)
        sec.left_margin = Pt(margin_pt + 10)
        sec.right_margin = Pt(margin_pt + 10)

    def _heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_pt + 2)
        p.paragraph_format.space_after = Pt(space_pt)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(body_pt + 1)
        pPr = p._p.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, _qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, _qn("w:bottom"))
        bottom.set(_qn("w:val"), "single")
        bottom.set(_qn("w:sz"), "4")
        bottom.set(_qn("w:space"), "1")
        bottom.set(_qn("w:color"), "auto")

    def _bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_pt)
        run = p.add_run(text)
        run.font.size = Pt(body_pt)

    def _spaced_para() -> object:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_pt)
        p.paragraph_format.space_after = Pt(0)
        return p

    # Header
    if resume.name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_pt)
        run = p.add_run(resume.name)
        run.bold = True
        run.font.size = Pt(name_pt)

    if resume.contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_pt)
        p.add_run(resume.contact).font.size = Pt(body_pt - 0.5)

    if resume.summary:
        _heading("Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_pt)
        p.add_run(resume.summary).font.size = Pt(body_pt)

    if resume.experience:
        _heading("Experience")
        for exp in resume.experience:
            p = _spaced_para()
            r = p.add_run(exp.title)
            r.bold = True
            r.font.size = Pt(body_pt)
            p.add_run(f"  ·  {exp.company}").font.size = Pt(body_pt)
            if exp.dates:
                p.add_run(f"  ·  {exp.dates}").font.size = Pt(body_pt - 0.5)
            for bullet in exp.bullets:
                _bullet(bullet)

    if resume.projects:
        _heading("Projects")
        for proj in resume.projects:
            p = _spaced_para()
            r = p.add_run(proj.name)
            r.bold = True
            r.font.size = Pt(body_pt)
            if proj.tech_stack:
                p.add_run(f"  ({', '.join(proj.tech_stack[:4])})").font.size = Pt(body_pt - 0.5)
            for bullet in proj.bullets:
                _bullet(bullet)

    if resume.skills:
        _heading("Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_pt)
        p.add_run(", ".join(resume.skills)).font.size = Pt(body_pt)

    if resume.education:
        _heading("Education")
        for edu in resume.education:
            parts = [edu.institution]
            if edu.degree:
                parts.append(edu.degree)
            if edu.field:
                parts.append(edu.field)
            if edu.year:
                parts.append(edu.year)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(space_pt)
            p.add_run(" | ".join(parts)).font.size = Pt(body_pt)

    doc.save(str(output_path))


# ── DOCX writing (tailored output) ───────────────────────────────────────────

def write_docx(original_path: str | Path, tailored: TailoredResume, output_path: str | Path) -> None:
    """Write tailored resume. Builds from scratch if base came from PDF, otherwise patches template."""
    with open(original_path, "rb") as f:
        is_pdf_base = f.read(4) == b"%PDF"
    marker = Path(original_path).parent / ".from_pdf"

    if is_pdf_base or marker.exists():
        resume = ResumeDocument(
            name=tailored.name,
            contact=tailored.contact,
            summary=tailored.summary,
            experience=tailored.experience,
            skills=tailored.skills,
            education=tailored.education,
            projects=tailored.projects,
        )
        build_docx_from_resume(resume, output_path)
        return

    shutil.copy2(str(original_path), str(output_path))
    doc = Document(str(output_path))

    current_section: Optional[str] = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            current_section = _detect_section(text)
            continue
        if current_section == "summary" and tailored.summary:
            _replace_paragraph_text(para, tailored.summary)
            tailored.summary = None
        elif current_section == "skills" and tailored.skills:
            _replace_paragraph_text(para, ", ".join(tailored.skills))
            tailored.skills = []

    doc.save(str(output_path))
    _enforce_single_page(Path(output_path))


def _enforce_single_page(docx_path: Path) -> None:
    """Reduce font sizes and spacing until the DOCX fits in 1 page."""
    pages = _get_page_count(docx_path)
    if pages is None or pages <= 1:
        return

    console.print(f"[yellow]{pages} pages — compacting to 1 page...[/yellow]")

    # Load once and mutate in memory; only save when needed for page-count check
    doc = Document(str(docx_path))
    for _ in range(5):
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.space_before and pf.space_before.pt > 0:
                pf.space_before = Pt(max(0, pf.space_before.pt - 1))
            if pf.space_after and pf.space_after.pt > 0:
                pf.space_after = Pt(max(0, pf.space_after.pt - 1))
            for run in para.runs:
                current = run.font.size.pt if run.font.size else 11
                run.font.size = Pt(max(7.5, current - 0.5))
        for sec in doc.sections:
            sec.top_margin = Pt(max(18, sec.top_margin.pt - 6))
            sec.bottom_margin = Pt(max(18, sec.bottom_margin.pt - 6))

        doc.save(str(docx_path))
        pages = _get_page_count(docx_path)
        if pages and pages <= 1:
            console.print("[green]Compacted to 1 page.[/green]")
            return

    console.print("[yellow]Warning: Could not compact to 1 page — content may be too long.[/yellow]")


def _replace_paragraph_text(para, new_text: str) -> None:
    if not para.runs:
        para.text = new_text
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


# ── ATS format check ──────────────────────────────────────────────────────────

def check_ats_format(path: str | Path) -> list[str]:
    """Return ATS format warnings for a DOCX. Skips check for PDFs."""
    with open(path, "rb") as f:
        if f.read(4) == b"%PDF":
            return ["File is a PDF — ATS format check skipped until DOCX is generated."]
    try:
        doc = Document(str(path))
    except Exception:
        return []

    warnings: list[str] = []
    if doc.tables:
        warnings.append(
            f"Document contains {len(doc.tables)} table(s). "
            "Tables may break ATS parsing — consider using plain text columns."
        )

    ats_safe_fonts = {"Calibri", "Arial", "Times New Roman", "Helvetica", "Georgia", "Garamond", "Trebuchet MS", "Verdana"}
    bad_fonts = {run.font.name for para in doc.paragraphs for run in para.runs if run.font.name and run.font.name not in ats_safe_fonts}
    if bad_fonts:
        warnings.append(f"Non-standard fonts: {', '.join(bad_fonts)}. Use Calibri, Arial, or Times New Roman.")

    return warnings


# ── PDF export ────────────────────────────────────────────────────────────────

def export_pdf(docx_path: str | Path, output_dir: str | Path) -> Optional[Path]:
    """Export DOCX to PDF using LibreOffice headless."""
    if _SOFFICE_BIN:
        try:
            result = subprocess.run(
                [_SOFFICE_BIN, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                pdf_path = Path(output_dir) / (Path(docx_path).stem + ".pdf")
                if pdf_path.exists():
                    return pdf_path
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass

    try:
        import pypandoc
        pdf_path = Path(output_dir) / (Path(docx_path).stem + ".pdf")
        pypandoc.convert_file(str(docx_path), "pdf", outputfile=str(pdf_path))
        return pdf_path
    except Exception:
        pass

    console.print("[yellow]Warning: PDF export failed. Install LibreOffice for PDF output.[/yellow]")
    return None
